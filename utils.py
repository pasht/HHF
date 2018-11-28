import mimetypes
import os
import zipfile
from tempfile import TemporaryDirectory
import codecs
import docx
import pandas as pd
from pandas import ExcelFile
import savReaderWriter as spss
import spacy
import re

nlp = spacy.load('el_core_web_sm')

def walkdir(path):
    """
    Recursively yield DirEntry objects for given directory.
    :param path: the root directory to start processing from
    """
    # print('Proccessing directory :{0}'.format(path))
    for entry in os.scandir(path):
        if entry.is_file():
            yield entry
        elif entry.is_dir(follow_symlinks=False):
            yield from walkdir(entry.path)


def capitalize(match):
    """
    capitalize a regex group
    :param match: regexp group object
    :return: capitilised sting
    """
    return match.group(1).capitalize()


def findNER(text):
    """
    Use Spacy to find PERSON Entities in test
    :param text: text to search for entities
    :return: a set containing PERSON NERs
    """
    doc = nlp(text)
    s: set = set()
    for entity in doc.ents:
        #if entity.label_ == 'PERSON' and entity.text is not None:
        if entity.label_ in ('PERSON', 'ORG') and len(entity.text.split()) == 2 and entity.text is not None:
            s.add(entity.text)
    return s
def dispatch(path, mtype):

    try:
        if 'zip' in mtype:
            zf = zipfile.ZipFile(path, 'r')
            tmpdir = TemporaryDirectory(dir=os.path.dirname(path))
            zf.extractall(tmpdir.name)
            for entry in walkdir(tmpdir.name):
                mimetype = mimetypes.guess_type(entry.path)[0]
                dispatch(entry.path, mimetype)

        # print('{0} is a {1}'.format(path, mtype))
        if 'spreadsheet' in mtype or 'ms-excel' in mtype:
            openexcelfile(path)
        if 'word' in mtype:
            opendocxfile(path)
        if 'spss' in mtype:
            opensavfile(path)
        if 'text' in mtype or 'csv' in mtype:
            opentxtfile(path)
    except Exception as ex:
        # print('File : {0} has a mime type that is not supported !!!'.format(path))
        print(ex)
    # print filename and file type

def openexcelfile(file):
    '''
    Read and/or process MS Excel file
    :param file: excel data or path to file to open it
    :return:
    '''
    xl = pd.ExcelFile(file)
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        t: str = re.sub("[\S\n\t]+]", " ", df.to_csv())
        t: str = re.sub("[/\\n\~-]+", " ", t)
        t: str = re.sub("(\w+)", capitalize, t)
        t: str = re.sub(',', " o ", t)
        s: set = findNER(t)
        if len(s) != 0:
            #print(s)
            print(file)

def opendocxfile(file):
    '''
     Read and/or process MS Word files
    :param: docx data or path to file to open it:
    '''
    doc = docx.Document(file)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    # Now print the extracted text
    text: str = re.sub("[._\\t\\n]+","",''.join(text))
    s: set = findNER(text)
    # print(s)
    if len(s) != 0:
        print(file)

def opensavfile(file):
    '''
     Read and/or process SPSS files
    :param file: sav data or path to file to open
    :return:
    '''
    text = []
    with spss.SavReader(file, ioUtf8=True) as reader:
        for line in reader:
            text.append(' '.join(str(element) for element in line))
    t: str = re.sub("[\\d,.\S]+", "", ''.join(text)).strip
    s = set()
    if len(s) > 0:
        s: set = findNER(''.join(text))
        if len(s) != 0:
            print(file)

def opentxtfile(file):
    '''
     Read and/or process text and csv files
    :param file: txt/csv data or path to file to open
    :return:
    '''
    s: set = set()
    with codecs.open(file, 'r', 'utf-8') as f:
        data = f.read()
        s: set = findNER(data)
    # Print the file name if we found NERs
    if len(s) != 0:
        print(file)
        #print(s)